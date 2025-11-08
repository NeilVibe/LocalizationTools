#!/usr/bin/env python3
"""
Excel/Word processor with parsing capabilities:
- Parse Excel files with context blocks and custom headers
- Convert Word tables to Excel
- Perform text cleaning and row filtering
"""

import os
import stat
import datetime
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import re
import traceback

def convert_doc_to_docx(doc_path):
    """Converts .doc file to .docx using win32com"""
    ext = os.path.splitext(doc_path)[1].lower()
    if ext != ".doc":
        return doc_path

    doc_path = os.path.normpath(doc_path)
    
    if not os.path.exists(doc_path):
        print(f"File not found: {doc_path}")
        return None

    try:
        mode = os.stat(doc_path).st_mode
        if not mode & stat.S_IWRITE:
            os.chmod(doc_path, mode | stat.S_IWRITE)
    except Exception as e:
        print(f"Permission error: {e}")
        traceback.print_exc()

    try:
        try:
            import win32com.client
        except ImportError:
            print("Missing package: pip install pywin32")
            return None

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(FileName=doc_path, ConfirmConversions=False, ReadOnly=True)
        new_path = os.path.splitext(doc_path)[0] + ".docx"
        doc.SaveAs(new_path, FileFormat=16)
        doc.Close()
        word.Quit()
        return new_path
    except Exception as e:
        print(f"Conversion error: {e}")
        traceback.print_exc()
        return None

def filter_rows_by_header(df, column_index, header_text):
    """Filter rows where the specified column contains the header text"""
    if column_index >= len(df.columns):
        return pd.DataFrame(columns=df.columns)
    
    # Create a mask for rows that have the header text in the specified column
    mask = df.iloc[:, column_index].astype(str).str.lower().str.startswith(header_text.lower())
    
    # Return filtered DataFrame
    return df[mask]

def generate_output_filename(input_filepath, suffix):
    """Generate output filename with timestamp"""
    folder, filename = os.path.split(input_filepath)
    name, ext = os.path.splitext(filename)
    dt_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"{name}_{suffix}_{dt_str}.xlsx"
    if not folder:
        folder = os.getcwd()
    return os.path.join(folder, output_filename)

def select_columns_dialog(columns):
    """Let user select columns to keep"""
    top = tk.Toplevel()
    top.title("Select Columns to Keep")
    selected_vars = {}
    frame = tk.Frame(top)
    frame.pack(padx=10, pady=10)
    
    for col in columns:
        var = tk.BooleanVar(value=True)
        cb = tk.Checkbutton(frame, text=col, variable=var)
        cb.pack(anchor="w")
        selected_vars[col] = var

    def on_ok():
        top.destroy()

    ok_btn = tk.Button(top, text="OK", command=on_ok)
    ok_btn.pack(pady=5)
    top.grab_set()
    top.wait_window()

    selected = [col for col in columns if selected_vars[col].get()]
    if not selected:
        selected = columns
    return selected

def process_word_file(file_path):
    """Process Word file table extraction"""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".doc":
        new_file_path = convert_doc_to_docx(file_path)
        if new_file_path is None:
            return None
        file_path = new_file_path

    try:
        from docx import Document
    except ImportError:
        print("Missing package: pip install python-docx")
        traceback.print_exc()
        return None

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f"Word read error: {e}")
        traceback.print_exc()
        return None

    if len(doc.tables) == 0:
        print("No tables found in document")
        return None

    desired_headers = ["CharacterName", "Original KR", "PA English"]
    selected_table = None

    for table in doc.tables:
        if len(table.rows) == 0:
            continue
        header_row = [cell.text.strip() for cell in table.rows[0].cells]
        if set(desired_headers).issubset(set(header_row)):
            selected_table = table
            break

    if selected_table is None:
        print("No table with expected headers found")
        return None

    data = []
    for row in selected_table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        data.append(cells)

    header_row = data[0]
    try:
        df = pd.DataFrame(data[1:], columns=header_row)
    except Exception as e:
        print(f"DataFrame creation error: {e}")
        traceback.print_exc()
        return None

    answer = messagebox.askyesno("Post-process Columns", 
                                f"File: {os.path.basename(file_path)}\n"
                                "Select columns to keep?")
    if answer:
        final_columns = select_columns_dialog(df.columns.tolist())
        df_final = df[final_columns]
    else:
        df_final = df

    return df_final

def reorder_files_dialog(file_paths):
    """Dialog to reorder selected files before processing"""
    if not file_paths:
        return []
        
    top = tk.Toplevel()
    top.title("Reorder Files")
    top.geometry("500x400")
    
    frame = tk.Frame(top)
    frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    label = tk.Label(frame, text="Drag files to reorder processing sequence:")
    label.pack(anchor="w", pady=(0, 10))
    
    # Convert file paths to just filenames for display
    file_names = [os.path.basename(path) for path in file_paths]
    
    # Create listbox with files
    listbox = tk.Listbox(frame, selectmode=tk.SINGLE, height=15)
    for name in file_names:
        listbox.insert(tk.END, name)
    listbox.pack(fill=tk.BOTH, expand=True)
    
    # Add move up/down buttons
    btn_frame = tk.Frame(frame)
    btn_frame.pack(fill=tk.X, pady=10)
    
    result = {"ordered_paths": None}
    
    def move_up():
        selected = listbox.curselection()
        if not selected or selected[0] == 0:
            return
        idx = selected[0]
        text = listbox.get(idx)
        listbox.delete(idx)
        listbox.insert(idx-1, text)
        listbox.selection_set(idx-1)
        listbox.see(idx-1)
    
    def move_down():
        selected = listbox.curselection()
        if not selected or selected[0] == listbox.size()-1:
            return
        idx = selected[0]
        text = listbox.get(idx)
        listbox.delete(idx)
        listbox.insert(idx+1, text)
        listbox.selection_set(idx+1)
        listbox.see(idx+1)
    
    def confirm_order():
        # Map display order back to full paths
        ordered_names = [listbox.get(i) for i in range(listbox.size())]
        name_to_path = {os.path.basename(path): path for path in file_paths}
        result["ordered_paths"] = [name_to_path[name] for name in ordered_names]
        top.destroy()
    
    up_btn = tk.Button(btn_frame, text="Move Up", command=move_up)
    up_btn.pack(side=tk.LEFT, padx=5)
    
    down_btn = tk.Button(btn_frame, text="Move Down", command=move_down)
    down_btn.pack(side=tk.LEFT, padx=5)
    
    confirm_btn = tk.Button(frame, text="Confirm Order", command=confirm_order)
    confirm_btn.pack(pady=10)
    
    top.grab_set()
    top.wait_window()
    
    return result["ordered_paths"] or []

def convert_word_files_to_excel_queue():
    """Process multiple Word files with ordering control"""
    file_paths = filedialog.askopenfilenames(title="Select Word File(s)",
                                           filetypes=[("Word files", "*.docx *.doc")])
    if not file_paths:
        return
    
    # Let user reorder files if more than one is selected
    if len(file_paths) > 1:
        reorder = messagebox.askyesno("File Order", "Would you like to control the order of file processing?")
        if reorder:
            file_paths = reorder_files_dialog(file_paths)
            if not file_paths:  # User canceled
                return
    
    processed_dfs = []
    for file_path in file_paths:
        try:
            df_final = process_word_file(file_path)
            if df_final is not None and not df_final.empty:
                processed_dfs.append(df_final)
                print(f"Processed: {file_path}")
            else:
                print(f"Skipped (no data): {file_path}")
        except Exception as e:
            print(f"Error processing: {file_path}\nError: {e}")
            traceback.print_exc()

    if not processed_dfs:
        messagebox.showinfo("Processing Complete", "No valid data extracted")
        return

    try:
        full_df = pd.concat(processed_dfs, ignore_index=True)
    except Exception as e:
        print(f"Concatenation error: {e}")
        traceback.print_exc()
        return

    output_filepath = generate_output_filename(file_paths[0], "concatenated")
    try:
        full_df.to_excel(output_filepath, index=False)
        messagebox.showinfo("Success", f"Concatenated Excel file created:\n{output_filepath}")
    except Exception as e:
        print(f"Excel write error: {e}")
        traceback.print_exc()

def clean_text_before_colon(df, selected_columns):
    """Remove everything before colon and strip whitespace in selected columns, leave empty cells as empty string"""
    for col in selected_columns:
        if col in df.columns:
            df[col] = df[col].apply(
                lambda x: re.sub(r'^.*:\s*', '', str(x)).strip() if pd.notna(x) and str(x).strip() != "" else ""
            )
    return df


def drop_empty_rows(df, column_selection, drop_if_any):
    """Drop rows based on empty cell conditions"""
    if df.empty:
        return df
        
    if drop_if_any:
        # Drop rows where ANY selected column is empty
        for col in column_selection:
            df = df[~df[col].isna()]
    else:
        # Drop rows where ALL selected columns are empty
        mask = pd.Series(False, index=df.index)
        for col in column_selection:
            mask = mask | df[col].notna()
        df = df[mask]
    
    return df

def open_excel_post_process_gui():
    """Create sub-GUI for Excel post-processing"""
    file_path = filedialog.askopenfilename(title="Select Excel File",
                                          filetypes=[("Excel files", "*.xlsx *.xls")])
    if not file_path:
        return
        
    try:
        # Try to read the Excel file with header
        df = pd.read_excel(file_path)
        if df.empty:
            messagebox.showerror("Error", "Excel file is empty")
            return
    except Exception as e:
        messagebox.showerror("Error", f"Cannot read Excel file:\n{str(e)}")
        return
    
    # Create processing sub-GUI
    sub_win = tk.Toplevel()
    sub_win.title("Excel Post-Processing")
    sub_win.geometry("650x550")
    
    # Create notebook (tabbed interface)
    notebook = ttk.Notebook(sub_win)
    notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    # Tab 1: Clean before colon
    tab1 = ttk.Frame(notebook)
    notebook.add(tab1, text="Clean Text")
    
    # Tab 2: Custom header parsing
    tab2 = ttk.Frame(notebook)
    notebook.add(tab2, text="Custom Header Parse")
    
    # Tab 3: Drop empty rows
    tab3 = ttk.Frame(notebook)
    notebook.add(tab3, text="Drop Empty Rows")
    
    # --- Tab 1: Clean before colon ---
    tk.Label(tab1, text="Clean text before colon character and strip whitespace").pack(anchor="w", pady=(10, 5))
    tk.Label(tab1, text="Example: 'ABCDE : HelloHello ' → 'HelloHello'").pack(anchor="w", pady=(0, 15))
    
    # Column selection for cleaning
    clean_frame = tk.LabelFrame(tab1, text="Select columns to clean")
    clean_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    clean_vars = {}
    for col in df.columns:
        var = tk.BooleanVar(value=False)
        cb = tk.Checkbutton(clean_frame, text=col, variable=var)
        cb.pack(anchor="w")
        clean_vars[col] = var
    
    def process_clean():
        selected_cols = [col for col, var in clean_vars.items() if var.get()]
        if not selected_cols:
            messagebox.showwarning("Warning", "No columns selected")
            return
            
        df_processed = clean_text_before_colon(df.copy(), selected_cols)
        save_processed_file(df_processed, file_path, "cleaned")
    
    tk.Button(tab1, text="Process and Save", command=process_clean).pack(pady=10)
    
    # --- Tab 2: Custom header parsing ---
    tk.Label(tab2, text="Filter rows based on custom header text").pack(anchor="w", pady=(10, 5))
    tk.Label(tab2, text="Define multiple parsing rules to filter rows").pack(anchor="w", pady=(0, 15))
    
    # Frame for parsing rules
    parsing_rules_frame = tk.LabelFrame(tab2, text="Parsing Rules")
    parsing_rules_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    # Create scrollable frame for parsing rules
    canvas = tk.Canvas(parsing_rules_frame)
    scrollbar = tk.Scrollbar(parsing_rules_frame, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas)
    
    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(
            scrollregion=canvas.bbox("all")
        )
    )
    
    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)
    
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    
    # List to hold parsing rules
    parsing_rules = []
    
    # Function to add new parsing rule
    def add_parsing_rule():
        rule_frame = tk.Frame(scrollable_frame)
        rule_frame.pack(fill=tk.X, pady=5)
        
        # Rule number
        rule_num = len(parsing_rules) + 1
        tk.Label(rule_frame, text=f"Rule {rule_num}:").pack(side=tk.LEFT, padx=(0, 10))
        
        # Column selection dropdown
        tk.Label(rule_frame, text="Column:").pack(side=tk.LEFT, padx=(0, 5))
        col_var = tk.StringVar()
        col_dropdown = ttk.Combobox(rule_frame, textvariable=col_var, values=list(df.columns), width=15)
        col_dropdown.pack(side=tk.LEFT, padx=(0, 10))
        
        # Header text entry
        tk.Label(rule_frame, text="Header:").pack(side=tk.LEFT, padx=(0, 5))
        header_entry = tk.Entry(rule_frame, width=15)
        header_entry.pack(side=tk.LEFT, padx=(0, 10))
        
        # Remove button
        def remove_rule():
            rule_frame.destroy()
            parsing_rules.remove((col_var, header_entry, rule_frame))
            # Update rule numbers for remaining rules
            for i, (_, _, frame) in enumerate(parsing_rules):
                for widget in frame.winfo_children():
                    if isinstance(widget, tk.Label) and widget.cget("text").startswith("Rule "):
                        widget.config(text=f"Rule {i+1}:")
        
        remove_btn = tk.Button(rule_frame, text="Remove", command=remove_rule)
        remove_btn.pack(side=tk.LEFT, padx=(0, 5))
        
        # Add to rules list
        parsing_rules.append((col_var, header_entry, rule_frame))
    
    # Add initial rule
    add_parsing_rule()
    
    # Button to add additional rules
    add_rule_btn = tk.Button(tab2, text="Add Rule", command=add_parsing_rule)
    add_rule_btn.pack(pady=5)
    
    def process_header_parsing():
        # Check if any rules are defined
        if not parsing_rules:
            messagebox.showwarning("Warning", "No parsing rules defined")
            return
        
        # Start with a copy of the original DataFrame
        filtered_df = df.copy()
        
        # Apply each rule
        for col_var, header_entry, _ in parsing_rules:
            col_name = col_var.get()
            header_text = header_entry.get().strip()
            
            if not col_name or not header_text:
                messagebox.showwarning("Warning", "Column and header text must be specified for all rules")
                return
            
            # Find the column index
            try:
                col_idx = list(df.columns).index(col_name)
            except ValueError:
                messagebox.showerror("Error", f"Column '{col_name}' not found in DataFrame")
                return
            
            # Apply the filter
            filtered_df = filter_rows_by_header(filtered_df, col_idx, header_text)
            
            if filtered_df.empty:
                messagebox.showinfo("Info", f"No rows match the filter criteria for '{header_text}' in column '{col_name}'")
                return
        
        # Save the filtered DataFrame
        save_processed_file(filtered_df, file_path, "filtered_headers")
    
    tk.Button(tab2, text="Process and Save", command=process_header_parsing).pack(pady=10)
    
    # --- Tab 3: Drop empty rows ---
    tk.Label(tab3, text="Drop rows based on empty cell conditions").pack(anchor="w", pady=(10, 5))
    
    # Add radio buttons for drop mode
    drop_mode = tk.StringVar(value="any")
    
    drop_mode_frame = tk.Frame(tab3)
    drop_mode_frame.pack(fill=tk.X, pady=10)
    
    tk.Radiobutton(drop_mode_frame, text="Drop if ANY selected column is empty", 
                   value="any", variable=drop_mode).pack(anchor="w")
    tk.Radiobutton(drop_mode_frame, text="Drop if ALL selected columns are empty", 
                   value="all", variable=drop_mode).pack(anchor="w")
    
    # Column selection for drop condition
    drop_frame = tk.LabelFrame(tab3, text="Select columns to check for emptiness")
    drop_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    
    drop_vars = {}
    for col in df.columns:
        var = tk.BooleanVar(value=False)
        cb = tk.Checkbutton(drop_frame, text=col, variable=var)
        cb.pack(anchor="w")
        drop_vars[col] = var
    
    def process_drop():
        selected_cols = [col for col, var in drop_vars.items() if var.get()]
        if not selected_cols:
            messagebox.showwarning("Warning", "No columns selected")
            return
            
        df_processed = drop_empty_rows(df.copy(), selected_cols, drop_mode.get() == "any")
        save_processed_file(df_processed, file_path, "filtered")
    
    tk.Button(tab3, text="Process and Save", command=process_drop).pack(pady=10)
    
    # Function to save processed DataFrames
    def save_processed_file(processed_df, original_path, suffix):
        if processed_df.empty:
            messagebox.showwarning("Warning", "Processing resulted in empty DataFrame")
            return
            
        output_path = generate_output_filename(original_path, suffix)
        try:
            processed_df.to_excel(output_path, index=False)
            messagebox.showinfo("Success", f"Processed file saved as:\n{output_path}")
        except Exception as e:
            messagebox.showerror("Error", f"Failed to save file:\n{str(e)}")

def create_gui():
    """Main application GUI"""
    root = tk.Tk()
    root.title("File Processor")
    root.geometry("400x200")

    # Word to Excel conversion button
    word_btn = tk.Button(root, text="Convert Word File(s) to Excel", 
                         command=convert_word_files_to_excel_queue, width=30, height=2)
    word_btn.pack(pady=20)
    
    # Excel post-processing button (now includes context parsing)
    post_process_btn = tk.Button(root, text="Excel Processing Tools", 
                                command=open_excel_post_process_gui, width=30, height=2)
    post_process_btn.pack(pady=20)

    root.mainloop()

if __name__ == "__main__":
    create_gui()